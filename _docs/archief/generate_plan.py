import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

doc = docx.Document()

# Title
title = doc.add_heading('Bedrijfsplan: Kustlab', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Executive Summary
add_heading(doc, '1. Executive Summary', 1)
doc.add_paragraph(
    "Kustlab, opgericht door Nick, is een Digital Solutions Consultancy bureau gevestigd in de regio Brugge. "
    "De transitie is ingezet van een louter uitvoerende website-bouwer naar een strategische partner voor kleine servicebedrijven. "
    "Het hoofddoel is niet alleen online zichtbaarheid vergroten via performante websites, maar voornamelijk het optimaliseren "
    "van de bedrijfsprocessen. Dit gebeurt door het bouwen van digitale workflows en automatisaties (zoals intake processen en "
    "offerte-generatie) die tijd besparen en de administratieve last verlagen."
)

# Bedrijfsoverzicht
add_heading(doc, '2. Missie en Positionering', 1)
doc.add_paragraph(
    "Missie:\n"
    "Lokale servicebedrijven empoweren door hun dagelijkse knelpunten op te lossen met slimme digitale oplossingen, "
    "waarbij techniek altijd de dienaar van het bedrijfsproces is, niet andersom."
)
doc.add_paragraph(
    "Positionering:\n"
    "Kustlab positioneert zichzelf als een 'High-Value Digital Solutions Consultant'. Geen 'code-klopper', maar een "
    "meedenkende partner die een 'Land & Expand' strategie toepast: instappen met een helder aanbod (website) "
    "en doorgroeien naar complexe, waarde-toevoegende integraties (workflow automatisatie)."
)

# Doelgroep
add_heading(doc, '3. Doelgroep (Target Audience)', 1)
doc.add_paragraph(
    "De primaire focus ligt op:\n"
    "- Kleine servicebedrijven en lokale aannemers (bijv. loodgieters, elektriciens, consultants, etc.) in en rond Brugge.\n"
    "- Bedrijven met minstens een paar medewerkers waarbij de eigenaar te veel in de operaties en administratie zit.\n"
    "- Bedrijven die nood hebben aan een professionele online uitstraling (Google Bedrijfsprofiel, snelle website) én "
    "gestroomlijnde processen (Lead captatie, geautomatiseerde afspraken)."
)

# Producten en Diensten (Pakketten)
add_heading(doc, '4. Dienstenaanbod & Prijsmodel', 1)
doc.add_paragraph(
    "Het aanbod is gestructureerd in 3 duidelijke pakketten om upselling te faciliteren ("
    "verkoop van resultaat, niet een uurtarief):"
)
doc.add_paragraph(
    "1. Instap Pakket (Foundations):\n"
    "Een sterke online basis. Moderne, snelle website of landing page, optimalisatie van Google Bedrijfsprofiel "
    "en basis contactformulieren. Bedoeld om lokaal gevonden te worden en vertrouwen te wekken."
)
doc.add_paragraph(
    "2. Midden Pakket (Automated Growth):\n"
    "De website aangevuld met specifieke intake-automatisering. Bijvoorbeeld de integratie van Calendly voor afspraken, "
    "geavanceerde intake-vragenlijsten om leads te kwalificeren vooraleer er telefonisch contact is, en koppeling aan een basis CRM/Lead Tracker."
)
doc.add_paragraph(
    "3. Premium Pakket (Digital Transformation):\n"
    "Een complete end-to-end oplossing. Naast de website worden workflows gebouwd via Google Scripts of vergelijkbare tools "
    "(zoals geautomatiseerde offertes op basis van intake). Dit omvat een volledige digitale optimalisatie van het acquisitieproces."
)

# Operationele en Intake flows
add_heading(doc, '5. Operationele Processen & Klantreis', 1)
doc.add_paragraph(
    "Aquisitie en Intake:"
)
doc.add_paragraph(
    "- Fase 1: 30-minuten Discovery Call.\n"
    "Puur kwalificerend. Ligt de focus op de juiste fit? Is er pijnpunt dat we kunnen oplossen?"
)
doc.add_paragraph(
    "- Fase 2: Diepgaande Strategie Sessie.\n"
    "Hier gaan we de diepte in aan de hand van gestructureerde intake-vragenlijsten. We identificeren knelpunten "
    "in hun huidige proces (hoe komt een klant binnen, hoe wordt de offerte gemaakt, waar verliest men tijd)."
)
doc.add_paragraph(
    "Delivery en Opschaling (Land & Expand):\n"
    "Zodra een klant is aan boord gebracht voor een website, worden pijnpunten gecapteerd om workflow-oplossingen als upsell te presenteren."
)

# Marketing & Sales Strategie
add_heading(doc, '6. Go-To-Market en Sales Strategie', 1)
doc.add_paragraph(
    "- Korte Termijn (30-day Growth Plan): Actieve cold email outreach. De pijpleiding wordt gevuld door 6 high-potential "
    "leads in de regio Brugge persoonlijk te benaderen met opgemaakte lead campagnes en campagneschetsen."
)
doc.add_paragraph(
    "- Tools: Lead_tracker.csv wordt gebruikt om potentiële klanten en hun fase in de pijplijn op te volgen. "
    "De positionering is overal consequent aangepast (professionele copy, website, Google Business Profile) naar 'Digital Solutions Consultant'."
)
doc.add_paragraph(
    "- Netwerk: Opbouwen van autoriteit door use-cases (bijv. de loodgieter website of de automapping van offertes) "
    "om te vormen tot succesverhalen die vertrouwen opwekken in koude acquisitie."
)

# Verdienmodel en Financieel
add_heading(doc, '7. Monetisatie', 1)
doc.add_paragraph(
    "- Eenmalige Inkomsten: Projectfees voor het opzetten van websites, domein design en workflow integraties."
)
doc.add_paragraph(
    "- Wederkerende Inkomsten (MRR): Onderhoudingscontacten voor websites, hosting, support bij automatisaties, "
    "en strategische opvolging (process monitoring, aanpassingen aan flow)."
)

doc.save('Bedrijfsplan_Kustlab.docx')
print("Business plan saved successfully to Bedrijfsplan_Kustlab.docx!")
